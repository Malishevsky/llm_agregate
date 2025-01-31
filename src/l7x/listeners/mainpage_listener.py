####################################################################################################
import asyncio
import base64
import os
from asyncio import sleep, get_running_loop
from functools import partial
from io import BytesIO
from typing import Final

from docx import Document
from nicegui import ui, App
from nicegui.elements.button import Button
from nicegui.elements.textarea import Textarea
from nicegui.events import UploadEventArguments
from pydantic.dataclasses import dataclass
from starlette.requests import Request
from torch import float16 as _torch_float16
from transformers import pipeline, PreTrainedTokenizer, PreTrainedModel

from l7x.configs.constants import RECONGIZER_MIME_TYPES
from l7x.configs.settings import AppSettings
from l7x.services.recognize_service import PrivateRecognizeService
from l7x.utils.cmd_manager_utils import BaseCommand, _CmdGlobalContext, _CmdLocalContext, _CmdReturnValue
from l7x.utils.fastapi_utils import AppFastAPI

#####################################################################################################

@dataclass(kw_only=True)
class AudioData:
    name: str
    type: str
    file: bytes

#####################################################################################################

@dataclass(kw_only=True, frozen=True)
class LlmProcessCommand(BaseCommand):
    text: str
    language: str
    convert_to: str | None

    #####################################################################################################

    async def execute(
        self,
        *,
        global_context: _CmdGlobalContext,
        local_context: _CmdLocalContext,
    ) -> _CmdReturnValue:
        model: Final[PreTrainedModel] = global_context.model
        tokenizer: Final[PreTrainedTokenizer] = global_context.tokenizer
        prompts_per_language: Final = global_context.app_settings.prompts_per_language
        prompts: Final = prompts_per_language.get(self.language, prompts_per_language['base'])
        if prompts is None:
            global_context.logger.warning('Prompts not found')
            return

        # sum_system_prompt = "Please summarize the text, highlighting the main topic, key points, and supporting details. Ensure your response is concise, accurate, and easy to understand. Your respond must be in the same language as the original sentence."
        text = self.text.strip()

        llm_pipeline = pipeline(
            "text-generation",
            model=model,
            tokenizer=tokenizer,
            model_kwargs={'torch_dtype': _torch_float16},
            device_map='auto',
        )
        summary_prompts = prompts.get('summary', [])
        new_prompts = summary_prompts

        if self.convert_to is not None:
            new_prompts.append(prompts.get(self.convert_to, []))

        for sum_prompt in new_prompts:
            messages = [
                {"role": "system", "content": sum_prompt},
                {"role": "user", "content": text}
            ]

            outputs = llm_pipeline(
                messages,
                max_new_tokens=10000,
                do_sample=False,
                temperature=0,
                pad_token_id=tokenizer.eos_token_id,
            )
            out_text = outputs[0]["generated_text"][-1]['content']

            text = out_text.strip()
        return text

#####################################################################################################

async def _summarize(
    summ_btn: Button,
    app: App,
    language: str | None,
    audio_data: AudioData | None,
    recognizer_area: Textarea,
    summarized_area: Textarea,
    radio_value: str,
):
    if audio_data is None:
        ui.notify('File not found', position='top', type='negative')
        return

    recognizer_service: PrivateRecognizeService = app.recognize_service

    try:
        summ_btn.set_enabled(False)
        recognized_text = await asyncio.wait_for(
            fut=recognizer_service.recognize(
                file_name=audio_data.name,
                mime_type=audio_data.type,
                wav=audio_data.file,
                language=language,
            ),
            timeout=180,
        )
        recognizer_area.set_value(recognized_text)
        await sleep(0.5)

        llm_cmd: Final = LlmProcessCommand(
            text=recognized_text,
            language=language,
            convert_to=radio_value.lower() if radio_value.lower() in ('formal', 'informal') else None,
        )

        loop = get_running_loop()
        send_and_wait_result = partial(
            app.cmd_manager.send_and_wait_result,
            call_timeout_sec=180,
        )
        summary_text = await loop.run_in_executor(None, send_and_wait_result, llm_cmd)
        summarized_area.set_value(summary_text)
    except Exception as err:
        print(f'Error: {err}')
        ui.notify('Recognize error', type='negative', position='top')
        return
    finally:
        summ_btn.set_enabled(True)

#####################################################################################################

def _add_text_with_bold(doc, text, heading: str | None = None):
    if heading is not None:
        doc.add_heading(heading)
    para = doc.add_paragraph()
    parts = text.split("**")
    for i, part in enumerate(parts):
        if i % 2 == 0:
            para.add_run(part)
        else:
            run = para.add_run(part)
            run.bold = True

#####################################################################################################

async def _download_file(rec_text: str, sum_text: str, audio_file_name: str):
    doc = Document()
    _add_text_with_bold(doc, rec_text, heading='Recognized Text')
    _add_text_with_bold(doc, sum_text, heading='Summarized Text')
    io_buffer = BytesIO()
    doc.save(io_buffer)
    io_buffer.seek(0)
    base64_string = base64.b64encode(io_buffer.read()).decode('utf-8')
    doc_file_name = os.path.splitext(audio_file_name)[0] + '_summarized.docx'
    await ui.run_javascript(f'''
        return await (async () => {{
            try {{
                const saveBase64File = async (base64String, fileName) => {{
                    const byteCharacters = atob(base64String);
                    const byteNumbers = new Array(byteCharacters.length);
                    for (let i = 0; i < byteCharacters.length; i++) {{
                        byteNumbers[i] = byteCharacters.charCodeAt(i);
                    }}
                    const byteArray = new Uint8Array(byteNumbers);
                    const blob = new Blob([byteArray], {{ type: 'application/octet-stream' }});
                    const link = document.createElement('a');
                    link.href = window.URL.createObjectURL(blob);
                    link.download = fileName;
                    document.body.appendChild(link);
                    link.click();
                    document.body.removeChild(link);
                    window.URL.revokeObjectURL(link.href);
                }};
                await saveBase64File("{base64_string}", "{doc_file_name}");
                return {{ success: true }};
            }} catch (error) {{
                throw new Error(error);
            }}
        }})();
    ''', timeout=30)

#####################################################################################################

async def _clear_handler(rec_area: Textarea, summ_area: Textarea, download_btn: Button):
    rec_area.set_value('')
    summ_area.set_value('')
    download_btn.set_visibility(False)

#####################################################################################################

async def _show_mainpage(request: Request) -> None:
    ui.add_css('./static/style.css')
    app: Final = request.app
    settings: Final[AppSettings] = app.settings
    audio_file_data: AudioData | None = None
    language = 'ru'
    if settings.dark_mode:
        ui.add_head_html('''
        <style>
            .shadow-border {box-shadow: 0 0 5px rgba(7, 167, 248, 0.9);}
        </style>
        ''')

    #######################################################################################

    async def _save_audio(event_args: UploadEventArguments) -> None:
        nonlocal audio_file_data
        if event_args.type not in RECONGIZER_MIME_TYPES:
            ui.notify('Not supported file type', type='negative', position='top')
            audio_file_data = None
            return

        audio_file_data = AudioData(
            name=event_args.name,
            type=event_args.type,
            file=event_args.content.read(),
        )

    #######################################################################################

    with ui.column().classes('main-container'):
        ui.upload(
            label='Audio file',
            auto_upload=True,
            max_file_size=settings.max_upload_file_size_in_byte,
            max_files=1,
            on_upload=_save_audio,
        ).classes('shadow-border')
        with ui.column().classes('main-block shadow-border'):
            ui.label('Select text style postprocessing')
            transform_radio = ui.radio(['FORMAL', 'INFORMAL', 'OFF'], value='OFF').props('inline')
            rec_area = ui.textarea(
                label='Recognized text',
            ).props('outlined readonly').classes('shadow-border')
            summ_area = ui.textarea(
                label='Summary text',
            ).props('outlined readonly').classes('shadow-border')
            with ui.row().classes('btn-container'):
                sum_btn = ui.button(
                    'Summarize',
                    icon='summarize',
                    on_click=lambda e: _summarize(
                        summ_btn=e.sender,
                        language=language,
                        app=app,
                        audio_data=audio_file_data,
                        recognizer_area=rec_area,
                        summarized_area=summ_area,
                        radio_value=transform_radio.value,
                    )
                )
                ui.button(
                    'Clear',
                    icon='delete',
                    on_click=lambda _: _clear_handler(
                        rec_area=rec_area,
                        summ_area=summ_area,
                        download_btn=download_btn,
                    ),
                ).classes('clear-btn').bind_enabled_from(sum_btn)

                download_btn = ui.button(
                    'Download file',
                    icon='download',
                    on_click=lambda _: _download_file(text=summ_area.value, audio_file_name=audio_file_data.name)
                ).bind_visibility_from(
                    summ_area,
                    target_name='value',
                    backward=lambda x: x != '',
                ).classes('download-btn')

            spin = ui.spinner(
                type='ios',
                size='3em',
            ).bind_visibility_from(
                sum_btn,
                target_name='enabled',
                backward=lambda x: not x,
            )
            spin.set_visibility(False)

#####################################################################################################

def mainpage_listener_registrar(app: AppFastAPI, /) -> None:
    ui.page('/', response_timeout=30)(_show_mainpage)

#####################################################################################################

